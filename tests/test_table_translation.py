"""
Tests for the Markdown-round-trip table translation feature.

Covers:
  - ParagraphBlock and TableBlock dataclasses
  - DocxProcessor.extract_blocks()
  - DocxProcessor.process_docx_for_translation()
  - TranslationService._rows_to_markdown()
  - TranslationService._parse_markdown_table()
  - TranslationService.translate_table_grid() (mocked API)
  - TranslationPromptSpec with has_table_markers=True
  - FileOutputHandler.save_to_docx() with table_registry
  - FileOutputHandler.save_translation_output() table_registry forwarding
"""

from io import BytesIO
from pathlib import Path
from typing import List
from unittest.mock import MagicMock, patch

import pytest

from src.models.doc_block import ParagraphBlock, TableBlock
from src.models import ParagraphBlock as PBlockFromInit, TableBlock as TBlockFromInit
from src.processors.docx_processor import DocxProcessor
from src.services.prompts.translation import TranslationPromptSpec
from src.services.prompts import fragments as F
from src.output.file_output import FileOutputHandler


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_docx(paragraphs: List[str] = (), rows_list: List[List[List[str]]] = ()) -> BytesIO:
    """Create an in-memory DOCX with the given paragraphs and tables (in that order)."""
    from docx import Document
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    for table_rows in rows_list:
        if not table_rows:
            continue
        n_cols = max(len(r) for r in table_rows)
        tbl = doc.add_table(rows=len(table_rows), cols=n_cols)
        for r_idx, row in enumerate(table_rows):
            for c_idx, cell_text in enumerate(row):
                tbl.cell(r_idx, c_idx).text = cell_text
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _make_mixed_docx() -> BytesIO:
    """Para → Table → Para document."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("Before table")
    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "H1"
    tbl.cell(0, 1).text = "H2"
    tbl.cell(1, 0).text = "R1C1"
    tbl.cell(1, 1).text = "R1C2"
    doc.add_paragraph("After table")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# ParagraphBlock and TableBlock dataclasses
# ---------------------------------------------------------------------------

class TestDocBlockModels:
    def test_paragraph_block_stores_text(self):
        block = ParagraphBlock(text="Hello world")
        assert block.text == "Hello world"

    def test_table_block_stores_rows_and_placeholder(self):
        rows = [["A", "B"], ["C", "D"]]
        block = TableBlock(rows=rows, placeholder="[TABLE_1]")
        assert block.rows == rows
        assert block.placeholder == "[TABLE_1]"

    def test_doc_block_models_exported_from_init(self):
        assert PBlockFromInit is ParagraphBlock
        assert TBlockFromInit is TableBlock


# ---------------------------------------------------------------------------
# DocxProcessor.extract_blocks()
# ---------------------------------------------------------------------------

class TestExtractBlocks:
    def test_only_paragraphs_all_become_paragraph_blocks(self):
        buf = _make_docx(paragraphs=["First", "Second"])
        blocks = DocxProcessor.extract_blocks(buf)
        assert all(isinstance(b, ParagraphBlock) for b in blocks)
        texts = [b.text for b in blocks]
        assert "First" in texts
        assert "Second" in texts

    def test_only_table_becomes_table_block(self):
        buf = _make_docx(rows_list=[[["A", "B"], ["C", "D"]]])
        blocks = DocxProcessor.extract_blocks(buf)
        assert len(blocks) == 1
        assert isinstance(blocks[0], TableBlock)

    def test_table_block_has_correct_placeholder(self):
        buf = _make_docx(rows_list=[[["A", "B"]]])
        blocks = DocxProcessor.extract_blocks(buf)
        assert blocks[0].placeholder == "[TABLE_1]"

    def test_multiple_tables_get_numbered_placeholders(self):
        buf = _make_docx(rows_list=[[["A"]], [["B"]], [["C"]]])
        blocks = DocxProcessor.extract_blocks(buf)
        placeholders = [b.placeholder for b in blocks if isinstance(b, TableBlock)]
        assert placeholders == ["[TABLE_1]", "[TABLE_2]", "[TABLE_3]"]

    def test_mixed_doc_preserves_document_order(self):
        buf = _make_mixed_docx()
        blocks = DocxProcessor.extract_blocks(buf)
        assert isinstance(blocks[0], ParagraphBlock)
        assert isinstance(blocks[1], TableBlock)
        assert isinstance(blocks[2], ParagraphBlock)
        assert blocks[0].text == "Before table"
        assert blocks[2].text == "After table"

    def test_table_block_cell_content(self):
        buf = _make_mixed_docx()
        blocks = DocxProcessor.extract_blocks(buf)
        tbl = next(b for b in blocks if isinstance(b, TableBlock))
        assert tbl.rows[0] == ["H1", "H2"]
        assert tbl.rows[1] == ["R1C1", "R1C2"]


# ---------------------------------------------------------------------------
# DocxProcessor.process_docx_for_translation()
# ---------------------------------------------------------------------------

class TestProcessDocxForTranslation:
    def test_pure_para_doc_returns_empty_registry(self):
        buf = _make_docx(paragraphs=["Hello", "World"])
        pages, registry = DocxProcessor.process_docx_for_translation(buf)
        assert registry == {}
        combined = "\n".join(pages)
        assert "Hello" in combined
        assert "World" in combined

    def test_table_placeholder_appears_in_pages(self):
        buf = _make_mixed_docx()
        pages, registry = DocxProcessor.process_docx_for_translation(buf)
        combined = "\n".join(pages)
        assert "[TABLE_1]" in combined

    def test_table_registry_populated(self):
        buf = _make_mixed_docx()
        _, registry = DocxProcessor.process_docx_for_translation(buf)
        assert "[TABLE_1]" in registry
        assert registry["[TABLE_1]"][0] == ["H1", "H2"]

    def test_multiple_tables_all_in_registry(self):
        buf = _make_docx(rows_list=[[["A"]], [["B"]]])
        _, registry = DocxProcessor.process_docx_for_translation(buf)
        assert "[TABLE_1]" in registry
        assert "[TABLE_2]" in registry

    def test_surrounding_paragraphs_not_lost(self):
        buf = _make_mixed_docx()
        pages, _ = DocxProcessor.process_docx_for_translation(buf)
        combined = "\n".join(pages)
        assert "Before table" in combined
        assert "After table" in combined

    def test_empty_doc_returns_sentinel(self):
        from docx import Document
        buf = BytesIO()
        Document().save(buf)
        buf.seek(0)
        pages, registry = DocxProcessor.process_docx_for_translation(buf)
        assert registry == {}
        # Should return a non-crashing placeholder
        assert isinstance(pages, list)


# ---------------------------------------------------------------------------
# TranslationService._rows_to_markdown()
# ---------------------------------------------------------------------------

class TestRowsToMarkdown:
    def _rto_md(self, rows):
        from src.services.translation_service import TranslationService
        return TranslationService._rows_to_markdown(rows)

    def test_single_row(self):
        md = self._rto_md([["Col A", "Col B"]])
        assert "| Col A | Col B |" in md

    def test_separator_row_inserted_after_header(self):
        md = self._rto_md([["H1", "H2"], ["R1", "R2"]])
        lines = md.splitlines()
        assert "---" in lines[1]

    def test_data_row_appears_after_separator(self):
        md = self._rto_md([["H1", "H2"], ["R1", "R2"]])
        lines = md.splitlines()
        assert "R1" in lines[2]

    def test_unequal_row_lengths_padded(self):
        md = self._rto_md([["A", "B", "C"], ["X"]])
        lines = md.splitlines()
        # Second data row (after separator) should have 3 pipe-separated cells
        data_line = [l for l in lines if "X" in l][0]
        assert data_line.count("|") >= 4  # 3 cells = 4 pipes

    def test_empty_rows_returns_empty_string(self):
        from src.services.translation_service import TranslationService
        assert TranslationService._rows_to_markdown([]) == ""


# ---------------------------------------------------------------------------
# TranslationService._parse_markdown_table()
# ---------------------------------------------------------------------------

class TestParseMarkdownTable:
    def _parse(self, md):
        from src.services.translation_service import TranslationService
        return TranslationService._parse_markdown_table(md)

    def test_valid_markdown_returns_rows(self):
        md = "| A | B |\n|---|---|\n| C | D |"
        rows = self._parse(md)
        assert rows is not None
        assert rows[0] == ["A", "B"]
        assert rows[1] == ["C", "D"]

    def test_separator_row_is_excluded(self):
        md = "| A | B |\n|---|---|\n| C | D |"
        rows = self._parse(md)
        assert all("---" not in cell for row in rows for cell in row)

    def test_non_table_input_returns_none(self):
        rows = self._parse("This is just plain text with no pipes.")
        assert rows is None

    def test_empty_string_returns_none(self):
        assert self._parse("") is None

    def test_extra_whitespace_stripped(self):
        md = "|  X  |  Y  |\n|---|---|\n|  1  |  2  |"
        rows = self._parse(md)
        assert rows[0] == ["X", "Y"]
        assert rows[1] == ["1", "2"]


# ---------------------------------------------------------------------------
# TranslationService.translate_table_grid() — mocked API
# ---------------------------------------------------------------------------

class TestTranslateTableGrid:
    def _make_service(self):
        from src.services.translation_service import TranslationService
        svc = TranslationService.__new__(TranslationService)
        svc.custom_model = None
        svc.custom_temperature = None
        svc.custom_top_p = None
        svc.custom_max_tokens = None
        svc.kanbun = False
        svc.system_note = None
        svc.user_note = None
        mock_tracker = MagicMock()
        mock_tracker.usage_data = {"total_usage": {"total_tokens": 0, "total_cost": 0.0}}
        svc.token_tracker = mock_tracker
        return svc

    def _mock_response(self, content: str):
        """Build a minimal mock API response."""
        choice = MagicMock()
        choice.message.content = content
        resp = MagicMock()
        resp.choices = [choice]
        resp.usage = MagicMock(prompt_tokens=10, completion_tokens=20, total_tokens=30)
        return resp

    def test_success_path_returns_translated_grid(self):
        svc = self._make_service()
        translated_md = "| Translated A | Translated B |\n|---|---|\n| Translated C | Translated D |"
        with patch.object(svc, '_get_model', return_value='gpt-4o'), \
             patch.object(svc, '_call_translation_api',
                         return_value=self._mock_response(translated_md)), \
             patch.object(svc, '_record_response_usage'):
            result = svc.translate_table_grid(
                [["A", "B"], ["C", "D"]], "Chinese", "English"
            )
        assert result[0] == ["Translated A", "Translated B"]
        assert result[1] == ["Translated C", "Translated D"]

    def test_row_count_mismatch_returns_original(self):
        svc = self._make_service()
        # Response has 3 rows but original has 2
        bad_md = "| X |\n|---|\n| Y |\n| Z |"
        with patch.object(svc, '_get_model', return_value='gpt-4o'), \
             patch.object(svc, '_call_translation_api',
                         return_value=self._mock_response(bad_md)), \
             patch.object(svc, '_record_response_usage'):
            original = [["A"], ["B"]]
            result = svc.translate_table_grid(original, "Japanese", "English")
        assert result == original

    def test_unparseable_response_returns_original(self):
        svc = self._make_service()
        with patch.object(svc, '_get_model', return_value='gpt-4o'), \
             patch.object(svc, '_call_translation_api',
                         return_value=self._mock_response("Sorry I cannot do that.")), \
             patch.object(svc, '_record_response_usage'):
            original = [["A", "B"]]
            result = svc.translate_table_grid(original, "Korean", "English")
        assert result == original

    def test_empty_rows_returns_empty_without_api_call(self):
        svc = self._make_service()
        with patch.object(svc, '_call_translation_api') as mock_api:
            result = svc.translate_table_grid([], "Chinese", "English")
        mock_api.assert_not_called()
        assert result == []


# ---------------------------------------------------------------------------
# TranslationPromptSpec — has_table_markers flag
# ---------------------------------------------------------------------------

class TestTranslationPromptSpecTableMarkers:
    def test_table_marker_rule_in_system_prompt_when_flag_set(self):
        spec = TranslationPromptSpec("Chinese", "English", has_table_markers=True)
        assert F.TRANSLATION_TABLE_MARKER_RULE in spec.system_prompt()

    def test_table_marker_rule_in_user_prompt_when_flag_set(self):
        spec = TranslationPromptSpec("Chinese", "English", has_table_markers=True)
        assert F.TRANSLATION_TABLE_MARKER_RULE in spec.user_prompt()

    def test_table_marker_rule_absent_by_default(self):
        spec = TranslationPromptSpec("Chinese", "English")
        assert F.TRANSLATION_TABLE_MARKER_RULE not in spec.system_prompt()
        assert F.TRANSLATION_TABLE_MARKER_RULE not in spec.user_prompt()

    def test_table_marker_rule_and_numbered_rule_coexist(self):
        spec = TranslationPromptSpec("Chinese", "English",
                                     has_table_markers=True, has_numbered=True)
        sys = spec.system_prompt()
        assert F.TRANSLATION_TABLE_MARKER_RULE in sys
        assert F.TRANSLATION_NUMBERED_SYSTEM in sys


# ---------------------------------------------------------------------------
# FileOutputHandler.save_to_docx() — table_registry reinsertion
# ---------------------------------------------------------------------------

class TestSaveToDocxWithTableRegistry:
    def test_placeholder_replaced_by_word_table(self, tmp_path):
        """A [TABLE_1] paragraph becomes a docx Table object, not a Paragraph."""
        from docx import Document
        out = str(tmp_path / "out.docx")
        registry = {"[TABLE_1]": [["H1", "H2"], ["R1", "R2"]]}
        content = "Intro paragraph\n\n[TABLE_1]\n\nConcluding paragraph"
        FileOutputHandler.save_to_docx(content, out, table_registry=registry)

        doc = Document(out)
        assert len(doc.tables) == 1
        tbl = doc.tables[0]
        assert tbl.cell(0, 0).text == "H1"
        assert tbl.cell(1, 1).text == "R2"

    def test_surrounding_paragraphs_preserved(self, tmp_path):
        from docx import Document
        out = str(tmp_path / "out.docx")
        registry = {"[TABLE_1]": [["A", "B"]]}
        content = "Before\n\n[TABLE_1]\n\nAfter"
        FileOutputHandler.save_to_docx(content, out, table_registry=registry)

        doc = Document(out)
        para_texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "Before" in para_texts
        assert "After" in para_texts

    def test_unknown_placeholder_written_as_plain_text(self, tmp_path):
        """A [TABLE_N] token with no matching registry entry stays as text."""
        from docx import Document
        out = str(tmp_path / "out.docx")
        # Pass no registry → placeholder treated as normal paragraph
        content = "Some text\n\n[TABLE_1]\n\nMore text"
        FileOutputHandler.save_to_docx(content, out, table_registry=None)

        doc = Document(out)
        assert len(doc.tables) == 0
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "[TABLE_1]" in all_text

    def test_multiple_tables_all_inserted(self, tmp_path):
        from docx import Document
        out = str(tmp_path / "out.docx")
        registry = {
            "[TABLE_1]": [["A"]],
            "[TABLE_2]": [["B"]],
        }
        content = "Para\n\n[TABLE_1]\n\nMiddle\n\n[TABLE_2]\n\nEnd"
        FileOutputHandler.save_to_docx(content, out, table_registry=registry)

        doc = Document(out)
        assert len(doc.tables) == 2

    def test_no_table_registry_produces_normal_output(self, tmp_path):
        from docx import Document
        out = str(tmp_path / "out.docx")
        FileOutputHandler.save_to_docx("Only paragraphs here.", out)
        doc = Document(out)
        assert len(doc.tables) == 0


# ---------------------------------------------------------------------------
# FileOutputHandler.save_translation_output() — table_registry forwarding
# ---------------------------------------------------------------------------

class TestSaveTranslationOutputTableForwarding:
    def test_table_registry_forwarded_to_save_to_docx(self, tmp_path):
        """save_translation_output passes table_registry on to save_to_docx."""
        registry = {"[TABLE_1]": [["X", "Y"]]}
        content = "Hello\n\n[TABLE_1]"
        out = str(tmp_path / "output.docx")
        with patch.object(FileOutputHandler, 'save_to_docx') as mock_save:
            FileOutputHandler.save_translation_output(
                content, None, out, False,
                "Chinese", "English",
                table_registry=registry,
            )
        mock_save.assert_called_once()
        _, kwargs = mock_save.call_args[0], mock_save.call_args[1]
        assert mock_save.call_args.kwargs.get('table_registry') == registry or \
               mock_save.call_args.args[-1] == registry or \
               any(a == registry for a in mock_save.call_args.args)

    def test_none_table_registry_does_not_crash(self, tmp_path):
        out = str(tmp_path / "output.docx")
        # Should not raise
        with patch.object(FileOutputHandler, 'save_to_docx'):
            FileOutputHandler.save_translation_output(
                "Content", None, out, False, "Chinese", "English",
                table_registry=None,
            )

    def test_txt_output_ignores_table_registry(self, tmp_path):
        """For .txt output, table_registry is irrelevant and not forwarded."""
        out = str(tmp_path / "output.txt")
        with patch.object(FileOutputHandler, 'save_to_text_file') as mock_txt:
            FileOutputHandler.save_translation_output(
                "Content", None, out, False, "Chinese", "English",
                table_registry={"[TABLE_1]": [["A"]]},
            )
        mock_txt.assert_called_once()
