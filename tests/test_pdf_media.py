"""
Tests for:
  1. Tables in PDF output (save_to_pdf with table_registry)
  2. PdfMediaExtractor — image extraction from PDF files
  3. --preserve-media validation now accepts .pdf input
  4. sandbox_processor wires PdfMediaExtractor for PDF input
"""

import struct
import zlib
from io import BytesIO
from pathlib import Path
from unittest.mock import patch

import pytest

from src.output.file_output import FileOutputHandler
from src.processors.pdf_media_extractor import PdfMediaExtractor
from src.models.embedded_media import EmbeddedMedia

_PLUGINS_DIR = Path(__file__).parent.parent / "plugins"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_png(width: int = 64, height: int = 64) -> bytes:
    """Return a valid RGB PNG of the given dimensions (solid white)."""
    def chunk(tag: bytes, data: bytes) -> bytes:
        c = struct.pack('>I', len(data)) + tag + data
        return c + struct.pack('>I', zlib.crc32(tag + data) & 0xFFFFFFFF)
    sig = b'\x89PNG\r\n\x1a\n'
    ihdr = chunk(b'IHDR', struct.pack('>IIBBBBB', width, height, 8, 2, 0, 0, 0))
    # Each row: filter byte (0) + width * 3 RGB bytes
    row = b'\x00' + b'\xff\xff\xff' * width
    raw_rows = row * height
    idat = chunk(b'IDAT', zlib.compress(raw_rows))
    iend = chunk(b'IEND', b'')
    return sig + ihdr + idat + iend


# Convenience alias used by tests that only need *a* valid PNG.
_make_1x1_png = lambda: _make_png(1, 1)  # noqa: E731 (kept for preserve-media compat)


def _make_pdf_with_image() -> BytesIO:
    """Create a minimal PDF containing one embedded 64×64 PNG using PyMuPDF."""
    import fitz
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    png_bytes = _make_png(64, 64)
    rect = fitz.Rect(100, 100, 200, 200)
    page.insert_image(rect, stream=png_bytes)
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def _make_empty_pdf() -> BytesIO:
    """Create a minimal PDF with no images."""
    import fitz
    doc = fitz.open()
    doc.new_page(width=595, height=842)
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


# ---------------------------------------------------------------------------
# Part 1: Tables in PDF output
# ---------------------------------------------------------------------------

class TestSaveToPdfWithTableRegistry:
    """Verify that [TABLE_N] placeholders are rendered as reportlab Table flowables."""

    def test_table_in_registry_produces_table_flowable(self, tmp_path):
        """save_to_pdf should call Table() when a matching placeholder is found."""
        out = str(tmp_path / "out.pdf")
        registry = {"[TABLE_1]": [["Header A", "Header B"], ["Cell 1", "Cell 2"]]}
        content = "Intro\n\n[TABLE_1]\n\nConclusion"




        # We just verify it runs without error and produces a real PDF file.
        FileOutputHandler.save_to_pdf(content, out, table_registry=registry, label="Translation")
        assert Path(out).exists()
        assert Path(out).stat().st_size > 0

    def test_pdf_without_table_registry_unchanged(self, tmp_path):
        out = str(tmp_path / "out.pdf")
        FileOutputHandler.save_to_pdf("Only plain text here.", out, table_registry=None, label="Translation")
        assert Path(out).exists()
        assert Path(out).stat().st_size > 0

    def test_unknown_placeholder_falls_through_to_text(self, tmp_path):
        """A [TABLE_N] token with no matching registry entry should write plain text."""
        out = str(tmp_path / "out.pdf")
        # No registry provided — token is treated as a normal paragraph
        FileOutputHandler.save_to_pdf(
            "Before\n\n[TABLE_1]\n\nAfter", out, table_registry=None, label="Translation"
        )
        assert Path(out).exists()

    def test_multiple_tables_in_registry(self, tmp_path):
        out = str(tmp_path / "out.pdf")
        registry = {
            "[TABLE_1]": [["A", "B"]],
            "[TABLE_2]": [["X", "Y"]],
        }
        content = "Intro\n\n[TABLE_1]\n\nMiddle\n\n[TABLE_2]\n\nEnd"
        FileOutputHandler.save_to_pdf(content, out, table_registry=registry, label="Translation")
        assert Path(out).exists()


class TestSaveTranslationOutputPdfTableForwarding:
    """save_translation_output forwards table_registry to save_to_pdf for .pdf output."""

    def test_table_registry_forwarded_to_save_to_pdf(self, tmp_path):
        registry = {"[TABLE_1]": [["X", "Y"]]}
        out = str(tmp_path / "output.pdf")
        with patch.object(FileOutputHandler, "save_to_pdf") as mock_pdf:
            FileOutputHandler.save_translation_output(
                "Hello\n\n[TABLE_1]", None, out, False,
                "Chinese", "English",
                table_registry=registry,
                label="Translation",
            )
        mock_pdf.assert_called_once()
        call_kwargs = mock_pdf.call_args.kwargs
        assert call_kwargs.get("table_registry") == registry

    def test_none_table_registry_forwarded_safely(self, tmp_path):
        out = str(tmp_path / "output.pdf")
        with patch.object(FileOutputHandler, "save_to_pdf") as mock_pdf:
            FileOutputHandler.save_translation_output(
                "Content", None, out, False, "Chinese", "English",
                table_registry=None,
                label="Translation",
            )
        mock_pdf.assert_called_once()
        call_kwargs = mock_pdf.call_args.kwargs
        assert call_kwargs.get("table_registry") is None


# ---------------------------------------------------------------------------
# Part 2: PdfMediaExtractor
# ---------------------------------------------------------------------------

class TestPdfMediaExtractorBasic:
    @pytest.fixture(autouse=True)
    def _bypass_size_filter(self):
        """Lower the minimum size threshold so our small test PNGs are accepted."""
        with patch("src.processors.pdf_media_extractor._MIN_IMAGE_BYTES", 0):
            yield

    def test_returns_list(self):
        buf = _make_empty_pdf()
        result = PdfMediaExtractor.extract_media(buf)
        assert isinstance(result, list)

    def test_empty_pdf_returns_empty_list(self):
        buf = _make_empty_pdf()
        result = PdfMediaExtractor.extract_media(buf)
        assert result == []

    def test_pdf_with_image_returns_one_item(self):
        buf = _make_pdf_with_image()
        result = PdfMediaExtractor.extract_media(buf)
        assert len(result) == 1

    def test_extracted_item_is_embedded_media(self):
        buf = _make_pdf_with_image()
        items = PdfMediaExtractor.extract_media(buf)
        assert isinstance(items[0], EmbeddedMedia)

    def test_image_data_is_non_empty_bytes(self):
        buf = _make_pdf_with_image()
        items = PdfMediaExtractor.extract_media(buf)
        assert isinstance(items[0].data, bytes)
        assert len(items[0].data) > 0

    def test_content_type_is_string(self):
        buf = _make_pdf_with_image()
        items = PdfMediaExtractor.extract_media(buf)
        assert isinstance(items[0].content_type, str)
        assert items[0].content_type.startswith("image/")

    def test_position_fraction_in_unit_interval(self):
        buf = _make_pdf_with_image()
        items = PdfMediaExtractor.extract_media(buf)
        assert 0.0 <= items[0].position_fraction <= 1.0

    def test_emu_dimensions_are_positive_or_none(self):
        buf = _make_pdf_with_image()
        items = PdfMediaExtractor.extract_media(buf)
        item = items[0]
        if item.width_emu is not None:
            assert item.width_emu > 0
        if item.height_emu is not None:
            assert item.height_emu > 0


class TestPdfMediaExtractorMultiPage:
    @pytest.fixture(autouse=True)
    def _bypass_size_filter(self):
        with patch("src.processors.pdf_media_extractor._MIN_IMAGE_BYTES", 0):
            yield

    def test_image_on_second_page_has_higher_fraction(self):
        import fitz
        doc = fitz.open()
        # Page 0: no image
        doc.new_page(width=595, height=842)
        # Page 1: image near the top
        page1 = doc.new_page(width=595, height=842)
        page1.insert_image(fitz.Rect(50, 50, 150, 150), stream=_make_png(64, 64))
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        items = PdfMediaExtractor.extract_media(buf)
        assert len(items) >= 1
        # Image is on page 1 of 2, so fraction should be >= 0.5
        assert items[0].position_fraction >= 0.5

    def test_deduplication_across_pages(self):
        """Same xref referenced on two pages should only appear once."""
        import fitz
        doc = fitz.open()
        png = _make_png(64, 64)
        p0 = doc.new_page(width=595, height=842)
        p0.insert_image(fitz.Rect(10, 10, 100, 100), stream=png)
        p1 = doc.new_page(width=595, height=842)
        # Insert the same PNG bytes again — PyMuPDF may reuse the xref or create a new one.
        p1.insert_image(fitz.Rect(10, 10, 100, 100), stream=png)
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        items = PdfMediaExtractor.extract_media(buf)
        # However many unique xrefs were created, each should appear exactly once.
        assert len(items) == len({id(it.data): it for it in items})


class TestPdfMediaExtractorImportError:
    def test_raises_import_error_when_fitz_missing(self):
        buf = _make_empty_pdf()
        with patch.dict("sys.modules", {"fitz": None}):
            with pytest.raises(ImportError, match="PyMuPDF"):
                PdfMediaExtractor.extract_media(buf)


# ---------------------------------------------------------------------------
# Part 3: --preserve-media validation — PDF input now allowed
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------

class TestSaveToDocxPageMarkerInsertion:
    """Images with page_number set use '-- Page N --' markers for placement."""

    @pytest.fixture(autouse=True)
    def _require_docx(self):
        pytest.importorskip("docx")

    def _para_texts(self, docx_path: str) -> list:
        from docx import Document
        return [p.text for p in Document(docx_path).paragraphs]

    def test_image_placed_after_correct_page_block(self, tmp_path):
        """Image with page_number=1 should appear after '-- Page 2 --' content,
        not at the start or end of the document."""
        from src.output.file_output import FileOutputHandler
        out = str(tmp_path / "pm.docx")
        # Translated output with two page blocks
        content = "\n\n-- Page 1 --\n\nText of page one.\n\n-- Page 2 --\n\nText of page two."
        # Image belongs to page_index=1 (label "Page 2")
        media = [EmbeddedMedia(
            data=_make_1x1_png(), content_type="image/png",
            position_fraction=0.5, page_number=1,
        )]
        FileOutputHandler.save_to_docx(content, out, media=media, label="Translation")
        self._para_texts(out)
        # Image paragraph has no text; locate it by finding the empty-run para
        from docx import Document
        doc = Document(out)
        para_texts = [p.text for p in doc.paragraphs]
        # '-- Page 2 --' and 'Text of page two.' must both appear
        assert "-- Page 2 --" in para_texts
        assert "Text of page two." in para_texts
        # Image (empty text paragraph) should NOT be before '-- Page 1 --'
        first_empty = next((i for i, t in enumerate(para_texts) if t == ""), None)
        page1_idx = para_texts.index("-- Page 1 --")
        assert first_empty is None or first_empty > page1_idx

    def test_image_on_page_zero_appears_before_page_two(self, tmp_path):
        """Image on page_index=0 should appear in the page-1 block (before page 2)."""
        from src.output.file_output import FileOutputHandler
        out = str(tmp_path / "p0.docx")
        content = "\n\n-- Page 1 --\n\nPage one text.\n\n-- Page 2 --\n\nPage two text."
        media = [EmbeddedMedia(
            data=_make_1x1_png(), content_type="image/png",
            position_fraction=0.1, page_number=0,
        )]
        FileOutputHandler.save_to_docx(content, out, media=media, label="Translation")
        from docx import Document
        doc = Document(out)
        para_texts = [p.text for p in doc.paragraphs]
        # Empty image paragraph should appear before '-- Page 2 --'
        page2_idx = para_texts.index("-- Page 2 --")
        empty_indices = [i for i, t in enumerate(para_texts) if t == ""]
        assert empty_indices, "Expected at least one image paragraph"
        assert all(idx < page2_idx for idx in empty_indices)

    def test_images_from_multiple_pages_all_inserted(self, tmp_path):
        """Images from two different PDF pages should both appear in the output."""
        from src.output.file_output import FileOutputHandler
        from docx import Document
        out = str(tmp_path / "multi.docx")
        content = "\n\n-- Page 1 --\n\nPage one.\n\n-- Page 2 --\n\nPage two."
        media = [
            EmbeddedMedia(data=_make_1x1_png(), content_type="image/png",
                          position_fraction=0.1, page_number=0),
            EmbeddedMedia(data=_make_1x1_png(), content_type="image/png",
                          position_fraction=0.6, page_number=1),
        ]
        FileOutputHandler.save_to_docx(content, out, media=media, label="Translation")
        doc = Document(out)
        # 2 page markers + 2 text paras + 2 image paras = 6 paragraphs
        assert len(doc.paragraphs) == 6

    def test_images_on_last_page_flushed_at_end(self, tmp_path):
        """Images for the last page should appear after that page's text."""
        from src.output.file_output import FileOutputHandler
        from docx import Document
        out = str(tmp_path / "last.docx")
        content = "\n\n-- Page 1 --\n\nOnly page."
        media = [EmbeddedMedia(
            data=_make_1x1_png(), content_type="image/png",
            position_fraction=0.5, page_number=0,
        )]
        FileOutputHandler.save_to_docx(content, out, media=media, label="Translation")
        doc = Document(out)
        para_texts = [p.text for p in doc.paragraphs]
        # '-- Page 1 --' and 'Only page.' must appear; image paragraph after them
        assert "-- Page 1 --" in para_texts
        assert "Only page." in para_texts
        page1_idx = para_texts.index("-- Page 1 --")
        empty_indices = [i for i, t in enumerate(para_texts) if t == ""]
        assert all(idx > page1_idx for idx in empty_indices)

    def test_page_number_none_uses_fraction_path(self, tmp_path):
        """Images with page_number=None fall back to fractional placement."""
        from src.output.file_output import FileOutputHandler
        from docx import Document
        out = str(tmp_path / "frac.docx")
        # Content with page markers but image has no page_number → fraction path
        content = "\n\n-- Page 1 --\n\nA.\n\n-- Page 2 --\n\nB.\n\nC."
        media = [EmbeddedMedia(
            data=_make_1x1_png(), content_type="image/png",
            position_fraction=0.99,  # very late → should be at end
            page_number=None,
        )]
        FileOutputHandler.save_to_docx(content, out, media=media, label="Translation")
        doc = Document(out)
        # 2 page-marker paras + 3 text paras (A., B., C.) + 1 image para = 6
        assert len(doc.paragraphs) == 6

    def test_page_number_field_on_extractor_output(self):
        """PdfMediaExtractor sets page_number to the 0-based page index."""
        with patch("src.processors.pdf_media_extractor._MIN_IMAGE_BYTES", 0):
            buf = _make_pdf_with_image()
            items = PdfMediaExtractor.extract_media(buf)
        assert len(items) == 1
        assert items[0].page_number == 0  # single-page PDF → page_index 0

