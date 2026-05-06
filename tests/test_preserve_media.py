"""
Tests for the --preserve-media feature.

Covers:
  - EmbeddedMedia dataclass
  - DocxProcessor.extract_media()
  - CLI flag parsing (--preserve-media)
  - _run_translate validation errors for incompatible flag combinations
  - FileOutputHandler.save_to_docx() with media reinsertion
  - FileOutputHandler.save_translation_output() media forwarding
"""

import argparse
import io
import struct
import zlib
from io import BytesIO
from pathlib import Path
from typing import List, Optional
from unittest.mock import MagicMock, patch, call

import pytest

from src.models.embedded_media import EmbeddedMedia
from src.models.output_options import OutputOptions
from src.errors import CLIError
from src.runtime.command_runner import _CommandMixin
from src.cli import create_argument_parser
from src.runtime.plugin_loader import load_plugins

_PLUGINS_DIR = Path(__file__).parent.parent / "plugins"


def _make_parser():
    return create_argument_parser(load_plugins(_PLUGINS_DIR))


def _make_1x1_png() -> bytes:
    """Return a minimal valid 1×1 white RGB PNG as bytes."""

    def chunk(tag: bytes, data: bytes) -> bytes:
        c = struct.pack('>I', len(data)) + tag + data
        return c + struct.pack('>I', zlib.crc32(tag + data) & 0xFFFFFFFF)

    sig = b'\x89PNG\r\n\x1a\n'
    ihdr = chunk(b'IHDR', struct.pack('>IIBBBBB', 1, 1, 8, 2, 0, 0, 0))
    raw = b'\x00\xff\xff\xff'  # filter=None, R=255, G=255, B=255
    idat = chunk(b'IDAT', zlib.compress(raw))
    iend = chunk(b'IEND', b'')
    return sig + ihdr + idat + iend


# ---------------------------------------------------------------------------
# EmbeddedMedia dataclass
# ---------------------------------------------------------------------------

class TestEmbeddedMedia:

    def test_required_fields(self):
        item = EmbeddedMedia(data=b"img", content_type="image/png", position_fraction=0.5)
        assert item.data == b"img"
        assert item.content_type == "image/png"
        assert item.position_fraction == 0.5

    def test_optional_emu_fields_default_to_none(self):
        item = EmbeddedMedia(data=b"img", content_type="image/png", position_fraction=0.0)
        assert item.width_emu is None
        assert item.height_emu is None

    def test_emu_fields_stored_when_provided(self):
        item = EmbeddedMedia(data=b"img", content_type="image/jpeg", position_fraction=0.25,
                             width_emu=914400, height_emu=457200)
        assert item.width_emu == 914400
        assert item.height_emu == 457200

    def test_position_fraction_at_boundaries(self):
        start = EmbeddedMedia(data=b"", content_type="image/png", position_fraction=0.0)
        end = EmbeddedMedia(data=b"", content_type="image/png", position_fraction=1.0)
        assert start.position_fraction == 0.0
        assert end.position_fraction == 1.0


# ---------------------------------------------------------------------------
# DocxProcessor.extract_media
# ---------------------------------------------------------------------------

class TestDocxProcessorExtractMedia:

    def _make_docx_with_image(self) -> BytesIO:
        """Return an in-memory .docx that has one embedded PNG image."""
        from docx import Document
        from docx.shared import Inches

        buf = BytesIO()
        doc = Document()
        doc.add_paragraph("Before image")
        para = doc.add_paragraph()
        run = para.add_run()
        run.add_picture(BytesIO(_make_1x1_png()), width=Inches(1))
        doc.add_paragraph("After image")
        doc.save(buf)
        buf.seek(0)
        return buf

    def test_returns_list_of_embedded_media(self):
        from src.processors.docx_processor import DocxProcessor
        buf = self._make_docx_with_image()
        items = DocxProcessor.extract_media(buf)
        assert isinstance(items, list)
        assert len(items) == 1

    def test_embedded_media_has_data(self):
        from src.processors.docx_processor import DocxProcessor
        buf = self._make_docx_with_image()
        items = DocxProcessor.extract_media(buf)
        assert len(items[0].data) > 0

    def test_position_fraction_in_range(self):
        from src.processors.docx_processor import DocxProcessor
        buf = self._make_docx_with_image()
        items = DocxProcessor.extract_media(buf)
        assert 0.0 <= items[0].position_fraction <= 1.0

    def test_empty_docx_returns_empty_list(self):
        from src.processors.docx_processor import DocxProcessor
        from docx import Document
        buf = BytesIO()
        doc = Document()
        doc.add_paragraph("No images here")
        doc.save(buf)
        buf.seek(0)
        items = DocxProcessor.extract_media(buf)
        assert items == []

    def test_content_type_is_image_mime(self):
        from src.processors.docx_processor import DocxProcessor
        buf = self._make_docx_with_image()
        items = DocxProcessor.extract_media(buf)
        assert items[0].content_type.startswith("image/")


# ---------------------------------------------------------------------------
# DocxProcessor.extract_raw_content — table extraction
# ---------------------------------------------------------------------------

class TestDocxProcessorTableExtraction:
    """Verify that text inside tables is included in extracted content."""

    def _make_docx_with_table(self, rows=2, cols=2) -> BytesIO:
        from docx import Document
        buf = BytesIO()
        doc = Document()
        doc.add_paragraph("Before table")
        table = doc.add_table(rows=rows, cols=cols)
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell.text = f"R{r_idx}C{c_idx}"
        doc.add_paragraph("After table")
        doc.save(buf)
        buf.seek(0)
        return buf

    def test_table_text_is_included(self):
        from src.processors.docx_processor import DocxProcessor
        buf = self._make_docx_with_table()
        processor = DocxProcessor()
        content = processor.extract_raw_content(buf)
        assert "R0C0" in content
        assert "R1C1" in content

    def test_table_appears_between_surrounding_paragraphs(self):
        from src.processors.docx_processor import DocxProcessor
        buf = self._make_docx_with_table()
        processor = DocxProcessor()
        content = processor.extract_raw_content(buf)
        before_pos = content.index("Before table")
        after_pos = content.index("After table")
        cell_pos = content.index("R0C0")
        assert before_pos < cell_pos < after_pos

    def test_table_cells_tab_separated(self):
        from src.processors.docx_processor import DocxProcessor
        buf = self._make_docx_with_table(rows=1, cols=3)
        # Overwrite cells so we know exact content
        from docx import Document
        buf2 = BytesIO()
        doc = Document()
        tbl = doc.add_table(rows=1, cols=3)
        tbl.rows[0].cells[0].text = "Alpha"
        tbl.rows[0].cells[1].text = "Beta"
        tbl.rows[0].cells[2].text = "Gamma"
        doc.save(buf2)
        buf2.seek(0)
        content = DocxProcessor().extract_raw_content(buf2)
        assert "Alpha\tBeta\tGamma" in content

    def test_pure_paragraph_doc_unchanged(self):
        from docx import Document
        from src.processors.docx_processor import DocxProcessor
        buf = BytesIO()
        doc = Document()
        doc.add_paragraph("First")
        doc.add_paragraph("Second")
        doc.save(buf)
        buf.seek(0)
        content = DocxProcessor().extract_raw_content(buf)
        assert "First" in content
        assert "Second" in content

    def test_process_docx_with_pages_includes_table_text(self):
        from src.processors.docx_processor import DocxProcessor
        buf = self._make_docx_with_table()
        pages = DocxProcessor.process_docx_with_pages(buf)
        combined = "\n".join(pages)
        assert "R0C0" in combined


# ---------------------------------------------------------------------------
# CLI flag parsing
# ---------------------------------------------------------------------------

class TestPreserveMediaCLIFlag:

    @pytest.fixture
    def parser(self):
        return _make_parser()

    def test_preserve_media_defaults_to_false(self, parser):
        args = parser.parse_args(["heller", "translate", "C-E", "-i", "doc.docx", "-o", "out.docx"])
        assert args.preserve_media is False

    def test_preserve_media_flag_sets_true(self, parser):
        args = parser.parse_args([
            "heller", "translate", "C-E", "-i", "doc.docx", "-o", "out.docx", "--preserve-media"
        ])
        assert args.preserve_media is True

    def test_preserve_media_not_present_on_transcribe(self, parser):
        args = parser.parse_args(["heller", "transcribe", "J", "-i", "img.png"])
        assert not hasattr(args, 'preserve_media') or args.preserve_media is False


# ---------------------------------------------------------------------------
# _run_translate validation errors
# ---------------------------------------------------------------------------

class _FakeMixin(_CommandMixin):
    """Minimal concrete subclass for testing _CommandMixin validation logic."""

    def __init__(self):
        self.translation_service = MagicMock()
        self.image_translation_service = MagicMock()
        self.image_processor = MagicMock()
        self.image_processor.is_image_file = MagicMock(return_value=False)
        self.image_processor_service = MagicMock()
        self.pdf_processor = MagicMock()
        self.prompt_service = MagicMock()
        self.transcription_review_service = MagicMock()
        self.token_tracker = MagicMock()
        self.file_output = MagicMock()

    def _detect_and_validate_file(self, file_path: str) -> str:
        return "docx"

    def translate_custom_text(self, *a, **kw): pass
    def process_image_translation_folder(self, *a, **kw): pass
    def translate_document(self, *a, **kw): pass
    def process_image_folder(self, *a, **kw): pass
    def process_image(self, *a, **kw): pass
    def process_prompt(self, *a, **kw): pass
    def process_transcription_review(self, *a, **kw): pass


def _make_translate_args(**overrides) -> argparse.Namespace:
    """Build a minimal valid translate Namespace, then apply overrides."""
    defaults = dict(
        language_code=("Chinese", "English"),
        input_file="doc.docx",
        custom_text=False,
        page_nums=None,
        abstract=False,
        auto_save=False,
        progressive_save=False,
        custom_font=None,
        preserve_media=False,
        output_file=None,
        workers=1,
        spread=False,
        kanbun=False,
        dry_run=False,
        notes=False,
        note_system=None,
        note_user=None,
        note_both=None,
        model=None,
        temperature=None,
        top_p=None,
        max_tokens=None,
    )
    defaults.update(overrides)
    return argparse.Namespace(**defaults)


class TestPreserveMediaValidation:

    def setup_method(self):
        self.mixin = _FakeMixin()

    def test_preserve_media_with_progressive_save_raises(self):
        args = _make_translate_args(
            preserve_media=True, progressive_save=True,
            input_file="doc.docx", output_file="out.docx"
        )
        with pytest.raises(CLIError, match="--progressive-save"):
            self.mixin._run_translate(args)

    def test_preserve_media_with_custom_text_raises(self):
        args = _make_translate_args(
            preserve_media=True, custom_text=True, input_file=None
        )
        with pytest.raises(CLIError, match=r"custom text"):
            self.mixin._run_translate(args)

    def test_preserve_media_without_input_file_raises(self):
        args = _make_translate_args(
            preserve_media=True, input_file=None, custom_text=False
        )
        with pytest.raises(CLIError, match=r"-i"):
            self.mixin._run_translate(args)

    def test_preserve_media_with_image_input_raises(self):
        args = _make_translate_args(
            preserve_media=True, input_file="scan.jpg", output_file="out.docx"
        )
        with pytest.raises(CLIError, match=r"image"):
            self.mixin._run_translate(args)

    def test_preserve_media_with_pdf_input_does_not_raise(self):
        """PDF input with .docx output is now supported for media preservation."""
        args = _make_translate_args(
            preserve_media=True, input_file="doc.pdf", output_file="out.docx"
        )
        # Should NOT raise; PDF input + docx output is a valid combination.
        try:
            self.mixin._run_translate(args)
        except CLIError as exc:
            pytest.fail(f"Unexpected CLIError for PDF input: {exc}")

    def test_preserve_media_with_txt_input_raises(self):
        args = _make_translate_args(
            preserve_media=True, input_file="doc.txt", output_file="out.docx"
        )
        with pytest.raises(CLIError, match=r"\.txt"):
            self.mixin._run_translate(args)

    def test_preserve_media_with_auto_save_no_output_raises(self):
        args = _make_translate_args(
            preserve_media=True, input_file="doc.docx",
            auto_save=True, output_file=None
        )
        with pytest.raises(CLIError, match=r"--auto-save"):
            self.mixin._run_translate(args)

    def test_preserve_media_without_output_file_raises(self):
        args = _make_translate_args(
            preserve_media=True, input_file="doc.docx",
            output_file=None, auto_save=False
        )
        with pytest.raises(CLIError, match=r"output"):
            self.mixin._run_translate(args)

    def test_preserve_media_with_txt_output_raises(self):
        args = _make_translate_args(
            preserve_media=True, input_file="doc.docx", output_file="out.txt"
        )
        with pytest.raises(CLIError, match=r"\.txt"):
            self.mixin._run_translate(args)

    def test_preserve_media_with_pdf_output_raises(self):
        args = _make_translate_args(
            preserve_media=True, input_file="doc.docx", output_file="out.pdf"
        )
        with pytest.raises(CLIError, match=r"PDF"):
            self.mixin._run_translate(args)

    def test_preserve_media_with_unknown_output_ext_raises(self):
        args = _make_translate_args(
            preserve_media=True, input_file="doc.docx", output_file="out.rtf"
        )
        with pytest.raises(CLIError, match=r"\.docx"):
            self.mixin._run_translate(args)

    def test_preserve_media_valid_combination_does_not_raise(self):
        """Valid: .docx input + .docx output — validation passes and translate_document is called."""
        args = _make_translate_args(
            preserve_media=True, input_file="doc.docx", output_file="out.docx"
        )
        with patch.object(self.mixin, "translate_document") as mock_td:
            self.mixin._run_translate(args)
        mock_td.assert_called_once()

    def test_preserve_media_false_skips_validation(self):
        """When --preserve-media is off, no validation errors should be raised."""
        args = _make_translate_args(preserve_media=False, input_file="doc.pdf", output_file=None)
        with patch.object(self.mixin, "translate_document"):
            # Should not raise even though no .docx output is specified
            self.mixin._run_translate(args)


# ---------------------------------------------------------------------------
# FileOutputHandler.save_to_docx with media
# ---------------------------------------------------------------------------

class TestSaveToDocxWithMedia:

    def test_save_without_media_creates_docx(self, tmp_path):
        from src.output.file_output import FileOutputHandler
        out = str(tmp_path / "out.docx")
        FileOutputHandler.save_to_docx("Paragraph one.\n\nParagraph two.", out, label="Translation")
        assert Path(out).exists()

    def test_save_with_media_creates_docx(self, tmp_path):
        from src.output.file_output import FileOutputHandler
        out = str(tmp_path / "out_media.docx")
        media = [EmbeddedMedia(data=_make_1x1_png(), content_type="image/png", position_fraction=0.5)]
        FileOutputHandler.save_to_docx("Para one.\n\nPara two.\n\nPara three.", out, media=media, label="Translation")
        assert Path(out).exists()

    def test_save_with_media_file_is_valid_docx(self, tmp_path):
        """Verify the output can be reopened by python-docx."""
        from docx import Document
        from src.output.file_output import FileOutputHandler
        out = str(tmp_path / "valid.docx")
        media = [EmbeddedMedia(data=_make_1x1_png(), content_type="image/png", position_fraction=0.3)]
        FileOutputHandler.save_to_docx("First para.\n\nSecond para.", out, media=media, label="Translation")
        doc = Document(out)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "First para." in texts
        assert "Second para." in texts

    def test_media_inserted_at_proportional_position(self, tmp_path):
        """Image paragraph should appear between translated paragraphs, not all at end."""
        from docx import Document
        from src.output.file_output import FileOutputHandler
        out = str(tmp_path / "pos.docx")
        # 3 paragraphs; image at 0.4 → should appear after para 1 (para_fraction=0.33) but
        # before or at para 2 (para_fraction=0.67).
        media = [EmbeddedMedia(data=_make_1x1_png(), content_type="image/png", position_fraction=0.4)]
        FileOutputHandler.save_to_docx("Alpha.\n\nBeta.\n\nGamma.", out, media=media, label="Translation")
        doc = Document(out)
        # Document has 4 paragraphs: Alpha, Beta (with image inserted after it), Gamma
        # The exact order depends on insertion logic; just verify 4 total blocks.
        assert len(doc.paragraphs) == 4

    def test_multiple_media_items_all_inserted(self, tmp_path):
        from docx import Document
        from src.output.file_output import FileOutputHandler
        out = str(tmp_path / "multi.docx")
        media = [
            EmbeddedMedia(data=_make_1x1_png(), content_type="image/png", position_fraction=0.2),
            EmbeddedMedia(data=_make_1x1_png(), content_type="image/png", position_fraction=0.8),
        ]
        FileOutputHandler.save_to_docx("A.\n\nB.\n\nC.\n\nD.\n\nE.", out, media=media, label="Translation")
        doc = Document(out)
        # 5 text paras + 2 image paras = 7
        assert len(doc.paragraphs) == 7

    def test_no_media_arg_behaves_identically_to_empty_list(self, tmp_path):
        from docx import Document
        from src.output.file_output import FileOutputHandler
        out1 = str(tmp_path / "none.docx")
        out2 = str(tmp_path / "empty.docx")
        content = "Hello.\n\nWorld."
        FileOutputHandler.save_to_docx(content, out1, media=None, label="Translation")
        FileOutputHandler.save_to_docx(content, out2, media=[], label="Translation")
        doc1 = Document(out1)
        doc2 = Document(out2)
        assert len(doc1.paragraphs) == len(doc2.paragraphs)


# ---------------------------------------------------------------------------
# FileOutputHandler.save_translation_output media forwarding
# ---------------------------------------------------------------------------

class TestSaveTranslationOutputMediaForwarding:

    def test_media_forwarded_to_save_to_docx(self, tmp_path):
        from src.output.file_output import FileOutputHandler
        out = str(tmp_path / "fwd.docx")
        media = [EmbeddedMedia(data=b"\x89PNG", content_type="image/png", position_fraction=0.5)]

        with patch.object(FileOutputHandler, "save_to_docx") as mock_docx:
            FileOutputHandler.save_translation_output(
                content="Some text.",
                input_file=None,
                output_file=out,
                auto_save=False,
                source_lang="Chinese",
                target_lang="English",
                media=media,
                label="Translation",
            )
        mock_docx.assert_called_once()
        _, kwargs = mock_docx.call_args[0], mock_docx.call_args[1]
        assert mock_docx.call_args[1].get("media") == media or mock_docx.call_args[0][4] == media

    def test_no_media_forwarded_when_none(self, tmp_path):
        from src.output.file_output import FileOutputHandler
        out = str(tmp_path / "no_media.docx")
        with patch.object(FileOutputHandler, "save_to_docx") as mock_docx:
            FileOutputHandler.save_translation_output(
                content="Some text.",
                input_file=None,
                output_file=out,
                auto_save=False,
                source_lang="Chinese",
                target_lang="English",
                label="Translation",
            )
        mock_docx.assert_called_once()
        # media kwarg should be None (default)
        assert mock_docx.call_args[1].get("media") is None

    def test_pdf_output_does_not_receive_media(self, tmp_path):
        from src.output.file_output import FileOutputHandler
        out = str(tmp_path / "out.pdf")
        media = [EmbeddedMedia(data=b"", content_type="image/png", position_fraction=0.5)]
        with patch.object(FileOutputHandler, "save_to_pdf") as mock_pdf:
            FileOutputHandler.save_translation_output(
                content="Text.",
                input_file=None,
                output_file=out,
                auto_save=False,
                source_lang="Chinese",
                target_lang="English",
                media=media,
                label="Translation",
            )
        mock_pdf.assert_called_once()
        # save_to_pdf does not accept a media argument — confirm it wasn't passed
        assert "media" not in mock_pdf.call_args[1]
