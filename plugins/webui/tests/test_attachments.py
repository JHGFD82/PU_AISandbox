"""Tests for plugins/webui/src/attachments.py — turning an uploaded document into chat-ready text.

Registered the same way conversation.py/app.py are (see plugin.py and
conftest.py), so this module is imported here via sys.modules under its
flat name rather than a normal dotted import.
"""

from __future__ import annotations

import json
import sys

import pytest

attachments = sys.modules["_pu_webui_attachments"]


def _write(tmp_path, name: str, content: bytes) -> str:
    path = tmp_path / name
    path.write_bytes(content)
    return str(path)


class TestUnsupportedAndMissingFiles:
    def test_unsupported_extension_raises(self, tmp_path):
        path = _write(tmp_path, "notes.exe", b"whatever")
        with pytest.raises(attachments.AttachmentError, match="isn't a supported file type"):
            attachments.extract_text(path, "notes.exe")

    def test_image_extension_is_not_supported(self, tmp_path):
        # Images are handled elsewhere (OCR/vision), not as text attachments.
        path = _write(tmp_path, "photo.png", b"\x89PNG\r\n\x1a\n")
        with pytest.raises(attachments.AttachmentError, match="isn't a supported file type"):
            attachments.extract_text(path, "photo.png")

    def test_missing_file_raises(self, tmp_path):
        with pytest.raises(attachments.AttachmentError, match="Couldn't read"):
            attachments.extract_text(str(tmp_path / "missing.txt"), "missing.txt")


class TestTxt:
    def test_extracts_plain_text(self, tmp_path):
        path = _write(tmp_path, "notes.txt", b"Hello from a text file.\n\nSecond paragraph.")
        doc = attachments.extract_text(path, "notes.txt")
        assert "Hello from a text file." in doc.text
        assert "Second paragraph." in doc.text
        assert doc.filename == "notes.txt"
        assert doc.char_count == len(doc.text)

    def test_empty_file_raises(self, tmp_path):
        path = _write(tmp_path, "empty.txt", b"")
        with pytest.raises(attachments.AttachmentError, match="No readable text"):
            attachments.extract_text(path, "empty.txt")


class TestMarkdown:
    def test_extracts_markdown(self, tmp_path):
        path = _write(tmp_path, "readme.md", b"# Title\n\nSome *markdown* content.")
        doc = attachments.extract_text(path, "readme.md")
        assert "Title" in doc.text
        assert "markdown" in doc.text


class TestJson:
    def test_extracts_json(self, tmp_path):
        path = _write(tmp_path, "data.json", json.dumps({"a": 1, "b": "text"}).encode())
        doc = attachments.extract_text(path, "data.json")
        assert "a" in doc.text and "1" in doc.text

    def test_invalid_json_raises(self, tmp_path):
        path = _write(tmp_path, "bad.json", b"{not valid json")
        with pytest.raises(attachments.AttachmentError, match="Couldn't read"):
            attachments.extract_text(path, "bad.json")


class TestDocx:
    def test_extracts_docx_paragraphs(self, tmp_path):
        from docx import Document

        doc_obj = Document()
        doc_obj.add_paragraph("First paragraph of the attached memo.")
        doc_obj.add_paragraph("Second paragraph with more detail.")
        path = tmp_path / "memo.docx"
        doc_obj.save(str(path))

        doc = attachments.extract_text(str(path), "memo.docx")
        assert "First paragraph of the attached memo." in doc.text
        assert "Second paragraph with more detail." in doc.text

    def test_corrupted_docx_raises(self, tmp_path):
        path = _write(tmp_path, "broken.docx", b"not a real docx file")
        with pytest.raises(attachments.AttachmentError, match="Couldn't read"):
            attachments.extract_text(path, "broken.docx")


class TestExcel:
    def test_extracts_excel_rows(self, tmp_path):
        import openpyxl

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws.append(["Name", "Score"])
        ws.append(["Ada", 95])
        path = tmp_path / "grades.xlsx"
        wb.save(str(path))

        doc = attachments.extract_text(str(path), "grades.xlsx")
        assert "Ada" in doc.text
        assert "Score" in doc.text


class TestPdf:
    def test_extracts_pdf_text(self, tmp_path):
        from reportlab.pdfgen import canvas

        path = tmp_path / "paper.pdf"
        c = canvas.Canvas(str(path))
        c.drawString(72, 700, "Hello PDF world, this is a test attachment.")
        c.save()

        doc = attachments.extract_text(str(path), "paper.pdf")
        assert "Hello PDF world" in doc.text


class TestSizeLimits:
    def test_upload_over_byte_cap_is_rejected_before_parsing(self, tmp_path, monkeypatch):
        monkeypatch.setattr(attachments, "MAX_UPLOAD_BYTES", 10)
        path = _write(tmp_path, "big.txt", b"x" * 100)
        with pytest.raises(attachments.AttachmentError, match="too large to attach"):
            attachments.extract_text(str(path), "big.txt")

    def test_extracted_text_over_char_cap_is_rejected_not_truncated(self, tmp_path, monkeypatch):
        monkeypatch.setattr(attachments, "MAX_ATTACHMENT_CHARS", 20)
        path = _write(tmp_path, "long.txt", b"This piece of text is definitely longer than twenty characters.")
        with pytest.raises(attachments.AttachmentError, match="too long to attach"):
            attachments.extract_text(str(path), "long.txt")

    def test_supported_extensions_excludes_images(self):
        assert ".png" not in attachments.SUPPORTED_EXTENSIONS
        assert ".pdf" in attachments.SUPPORTED_EXTENSIONS
        assert ".docx" in attachments.SUPPORTED_EXTENSIONS


class TestExtractedDocument:
    def test_char_count_matches_text_length(self, tmp_path):
        path = _write(tmp_path, "notes.txt", b"exactly this many characters")
        doc = attachments.extract_text(path, "notes.txt")
        assert doc.char_count == len(doc.text)
        assert isinstance(doc, attachments.ExtractedDocument)
