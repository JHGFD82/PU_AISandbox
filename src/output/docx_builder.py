"""Word document builder for translated content."""

import logging
from io import BytesIO
from pathlib import Path
from typing import List, Optional

from .font_resolver import get_docx_font
from ._output_utils import (
    _emit_message, _normalize_paragraphs, _extract_markdown_tables,
    _PAGE_MARKER_RE, save_to_text_file,
)
from ..settings import DEFAULT_FONT_SIZE

# Maximum width (EMU) an inserted image may occupy inside a Word document.
# Derived from US Letter (8.5") with 1" margins on each side: 6.5" × 914400 EMU/inch.
_MAX_IMAGE_WIDTH_EMU: int = 5_943_600


def _apply_docx_table_borders(table) -> None:  # type: ignore[no-untyped-def]
    """Apply thin (0.5 pt) black borders to all sides of a python-docx Table."""
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tbl_pr = table._tbl.tblPr
        tbl_borders = OxmlElement('w:tblBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')    # 4 × 1/8 pt = 0.5 pt
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), '000000')
            tbl_borders.append(el)
        tbl_pr.append(tbl_borders)
    except Exception as _be:
        logging.warning(f"Could not apply table borders: {_be}")


def save_to_docx(
    content: str,
    output_path: str,
    custom_font: Optional[str] = None,
    target_lang: Optional[str] = None,
    media: Optional[List] = None,
    table_registry: Optional[dict] = None,
    font_size: Optional[int] = None,
    *,
    label: str,
) -> None:
    """Save content to a Word document using python-docx.

    If *media* is provided (a list of :class:`~src.models.embedded_media.EmbeddedMedia`
    items), each image is reinserted at the proportionally equivalent position in
    the translated document.

    If *table_registry* is provided (a mapping from ``"[TABLE_N]"`` placeholder
    tokens to translated cell grids), any paragraph whose text is exactly a
    ``[TABLE_N]`` token is replaced by a proper Word ``Table`` object instead
    of being written as a paragraph.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Emu

        doc = Document()
        fs = font_size if font_size is not None else DEFAULT_FONT_SIZE

        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        if not custom_font and target_lang == 'English':
            font_name = 'Times New Roman'
            logging.debug(f"Using Times New Roman for English translation (target_lang={target_lang})")
        else:
            font_name = get_docx_font(custom_font)
            logging.debug(
                f"Using CJK font for Word: {font_name} "
                f"(custom_font={custom_font}, target_lang={target_lang})"
            )

        # --- Media routing -------------------------------------------------
        sorted_media = sorted(media, key=lambda m: m.position_fraction) if media else []

        pdf_media_by_page: dict[int, list] = {}
        frac_sorted: list = []
        for _m in sorted_media:
            if _m.page_number is not None:
                pdf_media_by_page.setdefault(_m.page_number, []).append(_m)
            else:
                frac_sorted.append(_m)

        frac_cursor = 0
        pending_pdf_images: list = []

        def _do_insert_image(item) -> None:  # type: ignore[no-untyped-def]
            """Write a single EmbeddedMedia item into the document."""
            try:
                img_para = doc.add_paragraph()
                img_run = img_para.add_run()
                display_width = Emu(
                    min(item.width_emu, _MAX_IMAGE_WIDTH_EMU)
                    if item.width_emu
                    else _MAX_IMAGE_WIDTH_EMU
                )
                img_data = BytesIO(item.data)
                try:
                    img_run.add_picture(img_data, width=display_width)
                except Exception:
                    try:
                        from PIL import Image as _PILImage
                        img_data.seek(0)
                        pil_img = _PILImage.open(img_data)
                        if pil_img.mode not in ('RGB', 'RGBA', 'L'):
                            pil_img = pil_img.convert('RGB')
                        png_buf = BytesIO()
                        pil_img.save(png_buf, format='PNG')
                        png_buf.seek(0)
                        img_run.add_picture(png_buf, width=display_width)
                        logging.debug(
                            f"Converted image at fraction "
                            f"{item.position_fraction:.3f} to PNG for insertion."
                        )
                    except Exception as conv_err:
                        logging.warning(
                            f"Could not insert image at fraction "
                            f"{item.position_fraction:.3f}: "
                            f"{type(conv_err).__name__}: {conv_err}"
                        )
                logging.debug(f"Inserted image at position_fraction={item.position_fraction:.3f}")
            except Exception as img_err:
                logging.warning(
                    f"Could not insert image at fraction {item.position_fraction:.3f}: "
                    f"{type(img_err).__name__}: {img_err}"
                )

        def _flush_pending_pdf_images() -> None:
            nonlocal pending_pdf_images
            for _item in pending_pdf_images:
                _do_insert_image(_item)
            pending_pdf_images = []

        def _insert_frac_images_up_to(para_fraction: float) -> None:
            nonlocal frac_cursor
            while (
                frac_cursor < len(frac_sorted)
                and frac_sorted[frac_cursor].position_fraction <= para_fraction
            ):
                _do_insert_image(frac_sorted[frac_cursor])
                frac_cursor += 1

        # Extract any inline Markdown tables and merge into table_registry.
        content, _md_reg = _extract_markdown_tables(content)
        if _md_reg:
            table_registry = dict(table_registry) if table_registry else {}
            table_registry.update(_md_reg)

        translated_paras = _normalize_paragraphs(content)
        total_paras = len(translated_paras) or 1

        for i, clean_text in enumerate(translated_paras, start=1):
            para_fraction = i / total_paras

            # --- Page-marker detection for PDF-source images ---------------
            if pdf_media_by_page:
                _pm = _PAGE_MARKER_RE.match(clean_text)
                if _pm:
                    _flush_pending_pdf_images()
                    _page_label = int(_pm.group(1))
                    pending_pdf_images = list(pdf_media_by_page.get(_page_label - 1, []))

            # --- Table placeholder: replace with a Word Table object ---
            if table_registry and clean_text.strip() in table_registry:
                rows = table_registry[clean_text.strip()]
                if rows:
                    try:
                        n_cols = max(len(row) for row in rows)
                        tbl = doc.add_table(rows=len(rows), cols=n_cols)
                        _apply_docx_table_borders(tbl)
                        for r_idx, row in enumerate(rows):
                            for c_idx, cell_text in enumerate(row):
                                if c_idx < n_cols:
                                    cell = tbl.cell(r_idx, c_idx)
                                    cell.text = cell_text
                                    for para in cell.paragraphs:
                                        for run in para.runs:
                                            run.font.name = font_name
                                            run.font.size = Pt(fs)
                        logging.debug(
                            f"Inserted Word table for '{clean_text.strip()}' "
                            f"({len(rows)} row(s) × {n_cols} col(s))"
                        )
                    except Exception as tbl_err:
                        logging.warning(
                            f"Could not insert table for '{clean_text.strip()}': {tbl_err}; "
                            "writing placeholder as plain text"
                        )
                        doc.add_paragraph(clean_text)
                _insert_frac_images_up_to(para_fraction)
                continue

            # --- Regular paragraph ---
            try:
                paragraph = doc.add_paragraph(clean_text)
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_after = Pt(12)
                paragraph_format.line_spacing = 1.5

                if paragraph.runs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run.font.size = Pt(fs)
                else:
                    run = paragraph.add_run(clean_text)
                    run.font.name = font_name
                    run.font.size = Pt(fs)

                logging.debug(f"Successfully added paragraph {i} with font {font_name}")
            except Exception as paragraph_error:
                logging.warning(f"Error processing paragraph {i} for Word document: {paragraph_error}")
                try:
                    doc.add_paragraph(clean_text)
                    logging.debug(f"Added paragraph {i} with basic formatting")
                except Exception as fallback_error:
                    logging.warning(f"Failed to add paragraph {i} to Word document: {fallback_error}")
                    _insert_frac_images_up_to(para_fraction)
                    continue

            _insert_frac_images_up_to(para_fraction)

        # Append any remaining images that fall past the last paragraph.
        _flush_pending_pdf_images()
        _insert_frac_images_up_to(1.0)

        if len(doc.paragraphs) > 0:
            doc.save(output_path)
            _emit_message(
                f"{label} saved to Word document: {Path(output_path).name}",
                level=logging.INFO,
                log_message=f'{label} saved to Word document: {output_path}',
                leading_newline=True,
            )
            if font_name != 'Times New Roman':
                _emit_message(f"Used font: {font_name}", level=logging.DEBUG)
        else:
            _emit_message(
                "Error: No content could be processed for Word document generation",
                level=logging.ERROR,
                log_message="No content could be processed for Word document generation",
            )
            _fallback_to_text(content, output_path, label)

    except ImportError:
        _emit_message(
            "Warning: python-docx not installed. To enable Word document export, install it with:",
            level=logging.WARNING,
            log_message='python-docx not installed. Falling back to text file.',
        )
        print("pip install python-docx")
        _emit_message("Saving as text file instead.", level=logging.WARNING)
        _fallback_to_text(content, output_path, label)
    except Exception as e:
        _emit_message(
            f"Error generating Word document: {e}",
            level=logging.ERROR,
            log_message=f'Error saving to Word document: {e}',
        )
        _emit_message(
            "Falling back to text file for reliable CJK character support...",
            level=logging.WARNING,
        )
        _fallback_to_text(content, output_path, label)


def _fallback_to_text(content: str, output_path: str, label: str) -> None:
    """Fallback to text output when Word document generation fails."""
    text_output_path = str(Path(output_path).with_suffix('.txt'))
    save_to_text_file(content, text_output_path, label)
