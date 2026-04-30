"""File output handler: writes translations to .txt, .pdf, or .docx with CJK font support."""

import logging
from io import BytesIO
from pathlib import Path
from typing import List, Optional
from datetime import datetime

from .font_resolver import get_docx_font, get_pdf_font

# PDF page margins (in points, 72 pts = 1 inch)
PDF_MARGINS = {
    'left': 72,
    'right': 72,
    'top': 72,
    'bottom': 18,
}

# Maximum width (EMU) an inserted image may occupy inside a Word document.
# Derived from US Letter (8.5") with 1" margins on each side: 6.5" × 914400 EMU/inch.
_MAX_IMAGE_WIDTH_EMU: int = 5_943_600



def generate_output_filename(input_file: str, source_lang: str, target_lang: str, extension: str = '.txt') -> str:
    """Generate an output filename based on input file and languages."""
    input_path = Path(input_file)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_name = f"{input_path.stem}_{source_lang}to{target_lang}_{timestamp}{extension}"
    return str(input_path.parent / output_name)


class FileOutputHandler:
    """Handles saving translations to various file formats."""

    @staticmethod
    def _emit_message(
        message: str,
        level: int = logging.INFO,
        log_message: Optional[str] = None,
        leading_newline: bool = False,
    ) -> None:
        """Emit a synchronized log + console message."""
        logging.log(level, log_message or message)
        prefix = "\n" if leading_newline else ""
        print(f"{prefix}{message}")

    @staticmethod
    def _normalize_paragraphs(content: str) -> list[str]:
        """Split content into normalized paragraphs for document output."""
        paragraphs: list[str] = []
        for paragraph in content.split('\n\n'):
            stripped = paragraph.strip()
            if stripped:
                paragraphs.append(stripped.replace('\n', ' '))
        return paragraphs

    @staticmethod
    def _resolve_output_path(
        input_file: Optional[str],
        output_file: Optional[str],
        auto_save: bool,
        source_lang: str,
        target_lang: str,
        default_extension: str = '.txt',
    ) -> Optional[str]:
        """Resolve the output path from explicit or auto-save settings."""
        if output_file:
            return output_file
        if auto_save and input_file:
            return generate_output_filename(input_file, source_lang, target_lang, default_extension)
        return None

    @staticmethod
    def _ensure_parent_directory(output_path: str) -> None:
        """Ensure the parent directory for output exists."""
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    @staticmethod
    def _fallback_to_text(content: str, output_path: str) -> None:
        """Fallback to text output when rich document generation fails."""
        text_output_path = str(Path(output_path).with_suffix('.txt'))
        FileOutputHandler.save_to_text_file(content, text_output_path)
    
    @staticmethod
    def save_to_text_file(content: str, output_path: str) -> None:
        """Save content to a text file."""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            FileOutputHandler._emit_message(
                f"Translation saved to: {Path(output_path).name}",
                level=logging.INFO,
                log_message=f'Translation saved to text file: {output_path}',
                leading_newline=True,
            )
        except (OSError, UnicodeError) as e:
            FileOutputHandler._emit_message(
                f"Error saving to text file: {e}",
                level=logging.ERROR,
            )

    @staticmethod
    def append_to_text_file(content: str, output_path: str) -> None:
        """Append content to a text file."""
        try:
            with open(output_path, 'a', encoding='utf-8') as f:
                f.write(content + '\n\n')
            FileOutputHandler._emit_message(
                f"Page appended to: {Path(output_path).name}",
                level=logging.INFO,
                log_message=f'Translation appended to text file: {output_path}',
            )
        except (OSError, UnicodeError) as e:
            FileOutputHandler._emit_message(
                f"Error appending to text file: {e}",
                level=logging.ERROR,
            )
    
    @staticmethod
    def save_to_pdf(content: str, output_path: str, custom_font: Optional[str] = None,
                    target_lang: Optional[str] = None,
                    table_registry: Optional[dict] = None) -> None:
        """Save content to a PDF file using reportlab.

        If *table_registry* is provided, ``[TABLE_N]`` placeholder paragraphs
        are rendered as reportlab ``Table`` flowables instead of plain text.
        """
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Flowable, Table, TableStyle
            from reportlab.lib import colors
            
            # Create PDF document
            doc = SimpleDocTemplate(
                output_path, 
                pagesize=letter,
                rightMargin=PDF_MARGINS['right'],
                leftMargin=PDF_MARGINS['left'],
                topMargin=PDF_MARGINS['top'],
                bottomMargin=PDF_MARGINS['bottom']
            )
            
            # Create story (content container)
            story: list[Flowable] = []
            styles = getSampleStyleSheet()
            
            # Configure font - use Times-Roman for English unless custom font specified
            if not custom_font and target_lang == 'English':
                font_name = 'Times-Roman'
                logging.debug(f"Using Times-Roman for English translation (target_lang={target_lang})")
            else:
                font_name = get_pdf_font(custom_font)
                logging.debug(f"Using CJK font: {font_name} (custom_font={custom_font}, target_lang={target_lang})")
            
            # Create paragraph style with proper font
            try:
                normal_style = ParagraphStyle(
                    'CJKNormal',
                    parent=styles['Normal'],
                    fontName=font_name,
                    fontSize=12,
                    leading=18,  # 1.5 leading (12pt * 1.5 = 18pt)
                    spaceAfter=12,
                    encoding='utf-8'
                )
                logging.debug(f"Created paragraph style with font: {font_name}")
            except (TypeError, ValueError, KeyError) as e:
                logging.warning(f"Failed to create custom style with font {font_name}: {e}")
                normal_style = styles['Normal']
                font_name = 'Times-Roman'  # Fallback to default
            
            fallback_style = ParagraphStyle(
                'FallbackCJK',
                parent=styles['Normal'],
                fontName='Times-Roman',
                fontSize=12,
                leading=18,
                spaceAfter=12
            )

            # Split content into paragraphs and add to story
            for i, clean_text in enumerate(FileOutputHandler._normalize_paragraphs(content), start=1):

                # --- Table placeholder: render as a reportlab Table flowable ---
                if table_registry and clean_text.strip() in table_registry:
                    rows = table_registry[clean_text.strip()]
                    if rows:
                        try:
                            # Build cell data — each cell wrapped in a Paragraph for wrapping
                            cell_data = [
                                [Paragraph(cell or '', normal_style) for cell in row]
                                for row in rows
                            ]
                            tbl = Table(cell_data, hAlign='LEFT')
                            tbl.setStyle(TableStyle([
                                ('GRID',       (0, 0), (-1, -1), 0.5, colors.black),
                                ('BACKGROUND', (0, 0), (-1, 0),  colors.lightgrey),
                                ('FONTNAME',   (0, 0), (-1, -1), font_name),
                                ('FONTSIZE',   (0, 0), (-1, -1), 10),
                                ('TOPPADDING',    (0, 0), (-1, -1), 4),
                                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                            ]))
                            story.append(tbl)
                            story.append(Spacer(1, 12))
                            logging.debug(
                                f"Inserted PDF table for '{clean_text.strip()}' "
                                f"({len(rows)} row(s))"
                            )
                            continue
                        except Exception as tbl_err:
                            logging.warning(
                                f"Could not render PDF table for '{clean_text.strip()}': {tbl_err}; "
                                "falling back to plain text"
                            )
                            # fall through to plain-paragraph rendering below

                try:
                    paragraph = Paragraph(clean_text, normal_style)
                    story.append(paragraph)
                    story.append(Spacer(1, 12))
                    logging.debug(f"Successfully added paragraph {i} with font {font_name}")
                except Exception as paragraph_error:
                    logging.warning(f"Error processing paragraph {i} with font {font_name}: {paragraph_error}")
                    try:
                        paragraph = Paragraph(clean_text, fallback_style)
                        story.append(paragraph)
                        story.append(Spacer(1, 12))
                        logging.debug(f"Used fallback font Times-Roman for paragraph {i}")
                    except Exception as fallback_error:
                        logging.warning(f"Fallback font also failed for paragraph {i}: {fallback_error}")
                        ascii_safe_text = clean_text.encode('ascii', 'ignore').decode('ascii')
                        if ascii_safe_text.strip():
                            paragraph = Paragraph(ascii_safe_text, styles['Normal'])
                            story.append(paragraph)
                            story.append(Spacer(1, 12))
                            logging.warning(f"Used ASCII-safe fallback for paragraph {i}")
                        else:
                            logging.warning(f"Paragraph {i} contained no ASCII-safe characters, skipping")
            
            # Build PDF
            if story:
                doc.build(story)
                FileOutputHandler._emit_message(
                    f"Translation saved to PDF: {Path(output_path).name}",
                    level=logging.INFO,
                    log_message=f'Translation saved to PDF file: {output_path}',
                    leading_newline=True,
                )
                if font_name != 'Times-Roman':
                    FileOutputHandler._emit_message(f"Used font: {font_name}", level=logging.DEBUG)
            else:
                FileOutputHandler._emit_message(
                    "Error: No content could be processed for PDF generation",
                    level=logging.ERROR,
                    log_message="No content could be processed for PDF generation",
                )
                FileOutputHandler._fallback_to_text(content, output_path)
            
        except ImportError:
            FileOutputHandler._emit_message(
                "Warning: reportlab not installed. Saving as text file instead.",
                level=logging.WARNING,
                log_message='reportlab not installed. Falling back to text file.',
            )
            FileOutputHandler._fallback_to_text(content, output_path)
        except Exception as e:
            FileOutputHandler._emit_message(
                f"Error generating PDF: {e}",
                level=logging.ERROR,
                log_message=f'Error saving to PDF: {e}',
            )
            FileOutputHandler._emit_message(
                "Falling back to text file for reliable CJK character support...",
                level=logging.WARNING,
            )
            FileOutputHandler._fallback_to_text(content, output_path)
    
    @staticmethod
    def save_to_docx(content: str, output_path: str, custom_font: Optional[str] = None,
                     target_lang: Optional[str] = None, media: Optional[List] = None,
                     table_registry: Optional[dict] = None) -> None:
        """Save content to a Word document using python-docx.

        If *media* is provided (a list of :class:`~src.models.embedded_media.EmbeddedMedia`
        items), each image is reinserted at the proportionally equivalent position in
        the translated document.

        If *table_registry* is provided (a mapping from ``"[TABLE_N]"`` placeholder
        tokens to translated cell grids), any paragraph whose text is exactly a
        ``[TABLE_N]`` token is replaced by a proper Word ``Table`` object instead
        of being written as a paragraph.
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, Emu

            # Create a new document
            doc = Document()

            # Set up margins (similar to PDF margins)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)

            # Get font name for CJK support
            if not custom_font and target_lang == 'English':
                font_name = 'Times New Roman'
                logging.debug(f"Using Times New Roman for English translation (target_lang={target_lang})")
            else:
                font_name = get_docx_font(custom_font)
                logging.debug(f"Using CJK font for Word: {font_name} (custom_font={custom_font}, target_lang={target_lang})")

            # Pre-sort media by position so we can sweep through in a single pass.
            sorted_media = sorted(media, key=lambda m: m.position_fraction) if media else []
            media_cursor = 0

            translated_paras = FileOutputHandler._normalize_paragraphs(content)
            total_paras = len(translated_paras) or 1

            def _insert_images_up_to(para_fraction: float) -> None:
                """Insert all media items whose position_fraction <= para_fraction."""
                nonlocal media_cursor
                while media_cursor < len(sorted_media) and sorted_media[media_cursor].position_fraction <= para_fraction:
                    item = sorted_media[media_cursor]
                    try:
                        img_para = doc.add_paragraph()
                        img_run = img_para.add_run()
                        # Cap width to the text-area; height scales automatically.
                        display_width = Emu(
                            min(item.width_emu, _MAX_IMAGE_WIDTH_EMU)
                            if item.width_emu
                            else _MAX_IMAGE_WIDTH_EMU
                        )
                        img_run.add_picture(BytesIO(item.data), width=display_width)
                        logging.debug(f"Inserted image at position_fraction={item.position_fraction:.3f}")
                    except Exception as img_err:
                        logging.warning(
                            f"Could not insert image at fraction {item.position_fraction:.3f}: "
                            f"{type(img_err).__name__}: {img_err}"
                        )
                    media_cursor += 1

            # Split content into paragraphs and add to document
            for i, clean_text in enumerate(translated_paras, start=1):
                para_fraction = i / total_paras

                # --- Table placeholder: replace with a Word Table object ---
                if table_registry and clean_text.strip() in table_registry:
                    rows = table_registry[clean_text.strip()]
                    if rows:
                        try:
                            n_cols = max(len(row) for row in rows)
                            tbl = doc.add_table(rows=len(rows), cols=n_cols)
                            for r_idx, row in enumerate(rows):
                                for c_idx, cell_text in enumerate(row):
                                    if c_idx < n_cols:
                                        cell = tbl.cell(r_idx, c_idx)
                                        cell.text = cell_text
                                        for para in cell.paragraphs:
                                            for run in para.runs:
                                                run.font.name = font_name
                                                run.font.size = Pt(12)
                            logging.debug(
                                f"Inserted Word table for '{clean_text.strip()}' "
                                f"({len(rows)} row(s) × {n_cols} col(s))"
                            )
                        except Exception as tbl_err:
                            logging.warning(
                                f"Could not insert table for '{clean_text.strip()}': {tbl_err}; "
                                "writing placeholder as plain text"
                            )
                            doc.add_paragraph(clean_text)
                    _insert_images_up_to(para_fraction)
                    continue

                # --- Regular paragraph ---
                try:
                    paragraph = doc.add_paragraph(clean_text)
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_after = Pt(12)
                    paragraph_format.line_spacing = 1.5

                    if paragraph.runs:
                        for run in paragraph.runs:
                            run.font.name = font_name
                            run.font.size = Pt(12)
                    else:
                        run = paragraph.add_run(clean_text)
                        run.font.name = font_name
                        run.font.size = Pt(12)

                    logging.debug(f"Successfully added paragraph {i} with font {font_name}")
                except Exception as paragraph_error:
                    logging.warning(f"Error processing paragraph {i} for Word document: {paragraph_error}")
                    try:
                        doc.add_paragraph(clean_text)
                        logging.debug(f"Added paragraph {i} with basic formatting")
                    except Exception as fallback_error:
                        logging.warning(f"Failed to add paragraph {i} to Word document: {fallback_error}")
                        para_fraction = i / total_paras
                        _insert_images_up_to(para_fraction)
                        continue

                _insert_images_up_to(para_fraction)

            # Append any remaining images that fall past the last paragraph.
            _insert_images_up_to(1.0)

            # Save the document
            if len(doc.paragraphs) > 0:
                doc.save(output_path)
                FileOutputHandler._emit_message(
                    f"Translation saved to Word document: {Path(output_path).name}",
                    level=logging.INFO,
                    log_message=f'Translation saved to Word document: {output_path}',
                    leading_newline=True,
                )
                if font_name != 'Times New Roman':
                    FileOutputHandler._emit_message(f"Used font: {font_name}", level=logging.DEBUG)
            else:
                FileOutputHandler._emit_message(
                    "Error: No content could be processed for Word document generation",
                    level=logging.ERROR,
                    log_message="No content could be processed for Word document generation",
                )
                FileOutputHandler._fallback_to_text(content, output_path)

        except ImportError:
            FileOutputHandler._emit_message(
                "Warning: python-docx not installed. To enable Word document export, install it with:",
                level=logging.WARNING,
                log_message='python-docx not installed. Falling back to text file.',
            )
            print("pip install python-docx")
            FileOutputHandler._emit_message("Saving as text file instead.", level=logging.WARNING)
            FileOutputHandler._fallback_to_text(content, output_path)
        except Exception as e:
            FileOutputHandler._emit_message(
                f"Error generating Word document: {e}",
                level=logging.ERROR,
                log_message=f'Error saving to Word document: {e}',
            )
            FileOutputHandler._emit_message(
                "Falling back to text file for reliable CJK character support...",
                level=logging.WARNING,
            )
            FileOutputHandler._fallback_to_text(content, output_path)
    
    @staticmethod
    def save_translation_output(content: str, input_file: Optional[str], output_file: Optional[str],
                              auto_save: bool, source_lang: str, target_lang: str,
                              custom_font: Optional[str] = None,
                              media: Optional[List] = None,
                              table_registry: Optional[dict] = None) -> None:
        """Save translation output to file based on user preferences."""
        if not content.strip():
            FileOutputHandler._emit_message("No content to save.", level=logging.INFO)
            return
        
        output_path = FileOutputHandler._resolve_output_path(
            input_file,
            output_file,
            auto_save,
            source_lang,
            target_lang,
            '.txt',
        )
        if not output_path:
            return
        
        FileOutputHandler._ensure_parent_directory(output_path)

        extension = Path(output_path).suffix.lower()
        if extension == '.pdf':
            FileOutputHandler.save_to_pdf(
                content, output_path, custom_font, target_lang,
                table_registry=table_registry,
            )
            return
        if extension == '.docx':
            FileOutputHandler.save_to_docx(
                content, output_path, custom_font, target_lang,
                media=media, table_registry=table_registry,
            )
            return

        if extension != '.txt':
            output_path = f"{output_path}.txt"
        FileOutputHandler.save_to_text_file(content, output_path)

    @staticmethod
    def save_page_progressively(content: str, input_file: Optional[str], output_file: Optional[str], 
                               auto_save: bool, source_lang: str, target_lang: str, is_first_page: bool = False,
                               custom_font: Optional[str] = None) -> Optional[str]:
        """Save a single page progressively to output file. Returns the output path."""
        _ = custom_font
        if not content.strip():
            FileOutputHandler._emit_message("No content to save.", level=logging.INFO)
            return None
        
        output_path = FileOutputHandler._resolve_output_path(
            input_file,
            output_file,
            auto_save,
            source_lang,
            target_lang,
            '.txt',
        )
        if not output_path:
            return None
        
        FileOutputHandler._ensure_parent_directory(output_path)
        
        # For progressive saving, we currently only support text files
        # PDF and Word document merging is complex and requires additional dependencies
        extension = Path(output_path).suffix.lower()
        if extension == '.pdf':
            FileOutputHandler._emit_message(
                "Note: Progressive saving for PDF format not yet supported. Using text format.",
                level=logging.INFO,
            )
            output_path = str(Path(output_path).with_suffix('.txt'))
        elif extension == '.docx':
            FileOutputHandler._emit_message(
                "Note: Progressive saving for Word document format not yet supported. Using text format.",
                level=logging.INFO,
            )
            output_path = str(Path(output_path).with_suffix('.txt'))
        
        # Default to text file
        if Path(output_path).suffix.lower() != '.txt':
            output_path = f"{output_path}.txt"
        
        # Save first page or append subsequent pages
        if is_first_page:
            FileOutputHandler.save_to_text_file(content, output_path)
        else:
            FileOutputHandler.append_to_text_file(content, output_path)
        
        return output_path
