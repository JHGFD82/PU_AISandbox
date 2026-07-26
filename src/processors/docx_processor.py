"""Word document processor: extracts text from .docx files into logical sections."""

import logging
from typing import List, BinaryIO, Union, TYPE_CHECKING

from .base_text_processor import BaseTextProcessor
from .constants import DEFAULT_PAGE_SIZE

if TYPE_CHECKING:
    from ..models.embedded_media import EmbeddedMedia
    from ..models.doc_block import ParagraphBlock, TableBlock


class DocxProcessor(BaseTextProcessor):
    """Extracts text and embedded media from Word (.docx) document files.

    Provides three extraction modes: plain text pages for translation,
    block-level extraction that keeps tables and paragraphs separate for
    round-trip DOCX output, and embedded image extraction for re-inserting
    images into the translated document.
    """
    
    def extract_raw_content(self, file_obj: BinaryIO) -> str:
        """Extract raw text content from a Word document, including table cells.

        Walks the document body in document order so that table content appears
        at its correct position relative to surrounding paragraphs.  Each table
        row is rendered as a tab-separated line; rows are separated by newlines.
        """
        try:
            from docx import Document
            from docx.oxml.ns import qn

            doc = Document(file_obj)
            blocks: List[str] = []

            for child in doc.element.body:
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

                if tag == 'p':
                    # Regular paragraph
                    text = ''.join(node.text or '' for node in child.iter(qn('w:t'))).strip()
                    if text:
                        blocks.append(text)

                elif tag == 'tbl':
                    # Table — collect rows, cells separated by tabs
                    rows: List[str] = []
                    for row in child.iter(qn('w:tr')):
                        cells: List[str] = []
                        for cell in row.iter(qn('w:tc')):
                            cell_text = ''.join(
                                node.text or '' for node in cell.iter(qn('w:t'))
                            ).strip()
                            cells.append(cell_text)
                        row_text = '\t'.join(cells)
                        if row_text.strip():
                            rows.append(row_text)
                    if rows:
                        blocks.append('\n'.join(rows))

            return '\n\n'.join(blocks) if blocks else ""

        except ImportError:
            raise ImportError(
                "python-docx is required to process Word documents. "
                "Install it with: pip install python-docx"
            ) from None
    
    @staticmethod
    def process_docx_with_pages(file_obj: BinaryIO, target_page_size: int = DEFAULT_PAGE_SIZE) -> List[str]:
        """Extract all text from a Word document and split it into translation-sized pages.

        Reads the document body in order (so table content appears at its
        correct position relative to surrounding paragraphs), then groups the
        resulting paragraphs into pages using the target character limit.
        Use this method when the output format is not ``.docx`` — for DOCX
        output, the translation plugin's own
        ``process_docx_for_translation`` (in
        ``plugins/translation/src/processors/docx_translation.py``) is used
        instead to preserve table structure.

        Args:
            file_obj: An open binary file object pointing to a ``.docx`` file.
            target_page_size: Approximate maximum characters per page.
                              Defaults to the project-wide page size setting.

        Returns:
            A list of page strings ready to be sent to the translation service,
            one string per logical page.
        """
        try:
            processor = DocxProcessor()
            content = processor.extract_raw_content(file_obj)
            
            if not content:
                logging.warning("No text content found in Word document")
                return [""]
            
            # Parse content into paragraphs and split into pages
            paragraphs = processor.parse_text_into_paragraphs(content)
            pages = processor.split_text_into_pages(paragraphs, target_page_size)
            
            logging.info(f"Split Word document into {len(pages)} logical pages")
            return pages
                
        except Exception as e:
            logging.error(f"Error processing Word document: {e}")
            raise Exception(f"Failed to process Word document: {e}") from e

    @staticmethod
    def extract_blocks(file_obj: BinaryIO) -> "List[Union[ParagraphBlock, TableBlock]]":
        """Read a Word document body and return its contents as an ordered list of typed blocks.

        Paragraphs are returned as ``ParagraphBlock`` items. Tables are
        returned as ``TableBlock`` items, each carrying a unique placeholder
        token (``[TABLE_1]``, ``[TABLE_2]``, etc.) that can be embedded in
        translation prompts. When the AI returns translated text, the
        placeholder positions mark exactly where each translated table should
        be reinserted in the output document. Document order is preserved so
        tables always appear at their correct position relative to surrounding
        paragraphs.

        Args:
            file_obj: An open binary file object pointing to a ``.docx`` file.

        Returns:
            A list of ``ParagraphBlock`` and ``TableBlock`` objects in document
            order.
        """
        from docx import Document
        from docx.oxml.ns import qn
        from ..models.doc_block import ParagraphBlock, TableBlock

        doc = Document(file_obj)
        blocks: List[Union[ParagraphBlock, TableBlock]] = []
        table_counter = 0

        for child in doc.element.body:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            if tag == 'p':
                text = ''.join(node.text or '' for node in child.iter(qn('w:t'))).strip()
                if text:
                    blocks.append(ParagraphBlock(text=text))

            elif tag == 'tbl':
                table_counter += 1
                placeholder = f"[TABLE_{table_counter}]"
                rows: List[List[str]] = []
                for row in child.iter(qn('w:tr')):
                    cells: List[str] = [
                        ''.join(node.text or '' for node in cell.iter(qn('w:t'))).strip()
                        for cell in row.iter(qn('w:tc'))
                    ]
                    if any(c for c in cells):
                        rows.append(cells)
                if rows:
                    blocks.append(TableBlock(rows=rows, placeholder=placeholder))

        return blocks

    @staticmethod
    def extract_media(file_obj: BinaryIO) -> "List[EmbeddedMedia]":
        """Pull all embedded images out of a Word document, recording each image's position in the document.

        Each returned ``EmbeddedMedia`` object carries a ``position_fraction``
        between 0.0 and 1.0 representing where in the document the image
        appeared relative to total paragraph count. This fraction is used when
        reinserting images into the translated document: because the translated
        text may have a different paragraph count, a proportional position
        keeps images near their original context rather than all piling up at
        the beginning or end.

        Args:
            file_obj: An open binary file object pointing to a ``.docx`` file.

        Returns:
            A list of ``EmbeddedMedia`` objects in document order, each
            containing the raw image bytes, content type, position fraction,
            and original dimensions. Returns an empty list if the document
            contains no embedded images.
        """
        from ..models.embedded_media import EmbeddedMedia
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(file_obj)
        paragraphs = doc.paragraphs
        total = len(paragraphs) or 1

        media_items: List[EmbeddedMedia] = []
        seen_r_ids: set = set()  # deduplicate reused image parts within the document

        for para_idx, para in enumerate(paragraphs):
            position_fraction = para_idx / total
            for run in para.runs:
                for drawing in run._element.iter(qn('w:drawing')):
                    for blip in drawing.iter(qn('a:blip')):
                        r_id = blip.get(qn('r:embed'))
                        if not r_id or r_id in seen_r_ids:
                            continue
                        try:
                            image_part = para.part.related_parts[r_id]
                        except KeyError:
                            logging.debug(f"Could not resolve image relationship '{r_id}'; skipping.")
                            continue

                        # Read EMU dimensions from the nearest wp:extent element.
                        width_emu: int | None = None
                        height_emu: int | None = None
                        for extent in drawing.iter(qn('wp:extent')):
                            try:
                                width_emu = int(extent.get('cx'))
                                height_emu = int(extent.get('cy'))
                            except (TypeError, ValueError):
                                pass
                            break  # only the first extent per drawing

                        seen_r_ids.add(r_id)
                        media_items.append(EmbeddedMedia(
                            data=image_part.blob,
                            content_type=image_part.content_type,
                            position_fraction=position_fraction,
                            width_emu=width_emu,
                            height_emu=height_emu,
                        ))

        logging.info(f"Extracted {len(media_items)} image(s) from Word document.")
        return media_items
